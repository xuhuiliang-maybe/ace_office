# coding=utf-8
import sys
import time

import pyExcelerator
from docx import Document

from modules.share_module.get_path import *

reload(sys)
sys.setdefaultencoding("utf-8")


class ExportExcel(object):
	def __int__(self, sheetname, head_title_list, field_name_list, data_obj_list, filename, path="tmp"):
		"""在指定的位置上生成一个excel 文件， 并将这个文件的路径作为返回值，供前台html 页面去下载
		:param sheetname:excel文件中sheet的名字
		:param head_title_list:列表类型的参数,列表每项是一个字符串。excel 文件中第一行的内容。
		:param field_name_list: 数据对应数据库字段名
		:param data_obj_list: 列表类型的参数,列表每项是一个列表，按照顺序显示每项的内容。  要导出的数据，
		:param filename: 生成的文件的前缀。
		:return:name，文件名
		"""
		self.sheetname = sheetname
		self.head_title_list = head_title_list
		self.field_name_list = field_name_list
		self.data_obj_list = data_obj_list
		self.filename = filename
		self.path = path

	def export(self):
		name = self.filename + "_" + time.strftime("%Y%m%d%H%M%S") + ".xls"
		tmp_path = get_media_sub_path(self.path)  # 临时文件夹路径
		path = os.path.join(tmp_path, name)  # 导出文件路径

		wb = pyExcelerator.Workbook()  # 创建一个pyExcelerator对象

		f1 = pyExcelerator.XFStyle()  # 标题样式
		f2 = pyExcelerator.XFStyle()  # 内容样式

		# 这里设置对齐方式
		align = pyExcelerator.Alignment()
		align.horz = pyExcelerator.Alignment.HORZ_CENTER
		align.vert = pyExcelerator.Alignment.VERT_CENTER
		f1.alignment = align
		f2.alignment = align

		# 这里设置字体大小
		f = pyExcelerator.Font()
		f.name = 'Arial'
		f.bold = True
		f.height = 240
		f1.font = f

		f = pyExcelerator.Font()
		f.name = 'Arial'
		f.height = 200
		f2.font = f
		try:

			if len(self.data_obj_list) == 0:
				return ""
			st = 0
			c = 0
			for i in xrange(len(self.data_obj_list)):
				if (i == 50000 * st):
					c = 0
					wf = wb.add_sheet(self.sheetname + str(st + 1))
					wf.col(1).width = 5000
					wf.col(2).width = 5000
					wf.col(3).width = 5000
					wf.col(4).width = 5000
					wf.col(5).width = 7000
					wf.col(6).width = 8000
					wf.col(7).width = 5000
					wf.col(9).width = 5000
					st += 1

					wf.write(0, 0, u"序号", f1)
					for t in xrange(len(self.head_title_list)):
						wf.write(0, t + 1, u"" + self.head_title_list[t] + "", f1)
				wf.write(i + 1, 0, str(i + 1), f2)
				for j in self.field_name_list:
					d = self.data_obj_list[i][j]
					if not d:
						d = "--"
					jj = self.field_name_list.index(j)
					try:
						wf.write(c + 1, jj + 1, d, f2)
					except:
						traceback.print_exc()
				c += 1
			wb.save(path)
			return path, name
		except:
			traceback.print_exc()
		return path, name


class ExportWord(object):
	def __int__(self, titles, content_list, filename, filepath):
		"""生成指定名称word文件，
		:param titles: 标题
		:param content_list:内容列表
		:param filename: 文件名
		:param filepath: 文件路径
		:return:
		"""
		self.titles = titles
		self.content_list = content_list
		self.filename = filename
		self.filepath = filepath

	def generate(self):
		document = Document()
		document.add_heading(self.titles, 0)

		for content in self.content_list:
			document.add_paragraph(content, style='ListBullet')

		document.add_page_break()
		document.save(r"%s" % (os.path.join(self.filepath, self.filename)))


class ExportTxt(object):
	def __int__(self, titles, content_list, filename, filepath):
		"""生成指定名称word文件，
		:param titles: 标题
		:param content_list:内容列表
		:param filename: 文件名
		:param filepath: 文件路径
		:return:
		"""
		self.titles = titles
		self.content_list = content_list
		self.filename = filename
		self.filepath = filepath

	def generate(self):
		with open(os.path.join(self.filepath, self.filename), "w") as new_txt:
			new_txt.write(self.titles + "\n\n")
			for one_row in self.content_list:
				new_txt.write(one_row + "\n")

		file_name_list = self.filename.split(".")
		word_file_name = file_name_list[0] + ".doc"
		os.rename(os.path.join(self.filepath, self.filename), os.path.join(self.filepath, word_file_name))


if __name__ == "__main__":
	word_obj = ExportWord()
	word_obj.titles = u"文章标题\n"
	word_obj.content_list = [u"时间：15455445", u"日期：45545"]
	word_obj.filename = "ds.doc"
	word_obj.filepath = "D:\\"
	word_obj.generate()

	txt_obj = ExportTxt()
	txt_obj.titles = u"文章标题"
	txt_obj.content_list = [u"时间：15455445", u"日期：45545"]
	txt_obj.filename = "ds.txt"
	txt_obj.filepath = "D:\\"
	txt_obj.generate()
